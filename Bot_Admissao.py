from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
import requests
from pathlib import Path
import glob
import pandas as pd
import time
from pypdf import PdfReader
from tabula import read_pdf
import PyPDF2
from PyPDF2 import PdfReader
import aspose.words as aw
import getpass
from docx import Document
from datetime import date

# DEF para chamado
def leiaInt(msg):
	ok = False
	valor = 0
	while True:
		NumChamado = str(input(msg))
		if NumChamado.isnumeric():
			valor = int(NumChamado)
			ok = True
		else:
			print('Erro! Digite um chamado valido!\n Digite novamente')
		if ok:
			break
	return valor

#Login
Login = str(input('Digite seu login: '))
#Senha
Senha = getpass.getpass('Digite sua senha:')
# Identificar o chamado
NumChamado = leiaInt('Digite o numero do chamado: ')
#Quantos arquivos existem:
quant_anexo = int(input('''Quantos arquivos tem em anexo dentro do Chamado: 
\n Escolha      \n[1]      \n[2]      \n[3]      \n[4]      \n[5]      \n[6] \n
\n *************************OBS************************* \n
    ABRIR DIRETO O ANEXO E VER QUANTIDADE DE ARQUIVOS
\n ***************************************************** \n''' ))
#Marcação do departamento
departamento = str(input("Qual departamento o Colaborador(a) pertence? \n DEPARTAMENTO: "))

#local do arquivo Chrome drive Versão 13.6
driver = webdriver.Chrome()
#Chamando o navegador
driver.get("https://moriah.qualitorsoftware.com/login.php")


#Variavel com o css selector do usuario
xpath_usu = "#cdusuario"
campo_usuario = driver.find_element("css selector", xpath_usu)
#Preenchendo o login
campo_usuario.send_keys(Login)

time.sleep(3)

#Criar campo para clicar fora, tem bloquei se não digitar a senha.
#Variavel com o css xpath da senha
css_selector_senha = "/html/body/div[3]/form/div[1]/div[2]/table/tbody/tr[2]/td/table/tbody/tr[1]/td/div[2]/table/tbody/tr[4]/td/span/input"
campo_senha = driver.find_element("xpath", css_selector_senha)
campo_senha.click()
campo_senha.click()
campo_senha.click()
#Preenchendo a senha
campo_senha.send_keys(Senha)


time.sleep(1)

#Clicando no botao de entrar
#variavel do botao entrar
bt_entrar = "/html/body/div[3]/form/div[1]/div[2]/table/tbody/tr[2]/td/table/tbody/tr[1]/td/div[2]/table/tbody/tr[9]/td/button"
#definindo elemento
clicar_entrar =  driver.find_element("xpath", bt_entrar)
#Click
clicar_entrar.click()

#abrindo uma nova pagina na web
driver.execute_script("window.open()")

time.sleep(15)
#trocando a aba
driver.switch_to.window(driver.window_handles[1])

# Carrega a nova aba
driver.get(f'https://moriah.qualitorsoftware.com/html/hd/hdchamado/cadastro_chamado.php?cdchamado={NumChamado}')
#Clicar em anexar
clicar_anexar =  driver.find_element("xpath", "/html/body/div[3]/form[1]/div[2]/div/div[1]/div[3]/div[3]").click()

if quant_anexo == 1:
	elemento_2_3_4 = "/html/body/div[3]/form[1]/div[2]/div/table/tbody/tr/td/div[3]/fieldset/div[2]/div/div/table/tbody/tr[2]/td[1]/span"
if quant_anexo == 2:
	elemento_2_3_4 = "/html/body/div[3]/form[1]/div[2]/div/table/tbody/tr/td/div[3]/fieldset/div[2]/div/div/table/tbody/tr[4]/td[1]/span"
if quant_anexo == 3:
	elemento_2_3_4 = "/html/body/div[3]/form[1]/div[2]/div/table/tbody/tr/td/div[3]/fieldset/div[2]/div/div/table/tbody/tr[6]/td[1]/span"
if quant_anexo == 4:
	elemento_2_3_4 = "/html/body/div[3]/form[1]/div[2]/div/table/tbody/tr/td/div[3]/fieldset/div[2]/div/div/table/tbody/tr[8]/td[1]/span"
if quant_anexo == 5:
	elemento_2_3_4 = "/html/body/div[3]/form[1]/div[2]/div/table/tbody/tr/td/div[3]/fieldset/div[2]/div/div/table/tbody/tr[10]/td[1]/span"
if quant_anexo == 6:
	elemento_2_3_4 = "/html/body/div[3]/form[1]/div[2]/div/table/tbody/tr/td/div[3]/fieldset/div[2]/div/div/table/tbody/tr[12]/td[1]/span"
if quant_anexo == 7:
	elemento_2_3_4 = "/html/body/div[3]/form[1]/div[2]/div/table/tbody/tr/td/div[3]/fieldset/div[2]/div/div/table/tbody/tr[14]/td[1]/span"
if quant_anexo == 8:
	elemento_2_3_4 = "/html/body/div[3]/form[1]/div[2]/div/table/tbody/tr/td/div[3]/fieldset/div[2]/div/div/table/tbody/tr[16]/td[1]/span"
if quant_anexo == 9:
	elemento_2_3_4 = "/html/body/div[3]/form[1]/div[2]/div/table/tbody/tr/td/div[3]/fieldset/div[2]/div/div/table/tbody/tr[18]/td[1]/span"
if quant_anexo == 10:
	elemento_2_3_4 = "/html/body/div[3]/form[1]/div[2]/div/table/tbody/tr/td/div[3]/fieldset/div[2]/div/div/table/tbody/tr[20]/td[1]/span"

time.sleep(5)
#Definindo quantidade de anexos
clicar_baixar_todos =  driver.find_element("xpath", elemento_2_3_4).click()

time.sleep(10)

driver.close()

# carregar arquivo de texto
doc = aw.Document(f"C://Users/thiago.asouza/Downloads/CRIACAO_ALTERACAO_DE_ACESSOS_(NOVO)_{NumChamado}.html")

# salvar texto como HTML
doc.save(f"{NumChamado}.txt")

with open(f"{NumChamado}.txt" , "r" , encoding="utf-8") as arquivo2:
    texto = arquivo2.readlines()

	
#Manipulando variavel do texto do arquivo txt
if "Nome" in texto[12]:
	Nome_texto = texto[13]
	nome_split = Nome_texto.split()
	cpf_texto = texto[23]
	cpf_split = cpf_texto.split()
	cpf = cpf_split[0]

if "Nome" in texto[13]:
	Nome_texto = texto[14]
	nome_split = Nome_texto.split()
	cpf_texto = texto[24]
	cpf_split = cpf_texto.split()
	cpf = cpf_split[0]
	
if "Nome" in texto[14]:#TINHA TESTADO POR AQUI
	Nome_texto = texto[15]
	nome_split = Nome_texto.split()
	cpf_texto = texto[25]
	cpf_split = cpf_texto.split()
	cpf = cpf_split[0]

	
if "Nome" in texto[15]:
	Nome_texto = texto[16]
	nome_split = Nome_texto.split()
	cpf_texto = texto[26]
	cpf_split = cpf_texto.split()
	cpf = cpf_split[0]


if "Nome" in texto[16]:
	Nome_texto = texto[17]
	nome_split = Nome_texto.split()
	cpf_texto = texto[27]
	cpf_split = cpf_texto.split()
	cpf = cpf_split[0]	
	
#Manipulando variavel LOGIN e NOME
if len(nome_split) == 2:
	Nome = (f"{nome_split[0]} {nome_split[1]}")
	login = (f"{nome_split[0]}.{nome_split[1]}")
if len(nome_split) == 3:
	Nome = (f"{nome_split[0]} {nome_split[1]} {nome_split[2]}")
	login = (f"{nome_split[0]}.{nome_split[2]}")
if len(nome_split) == 4:
	Nome = (f"{nome_split[0]} {nome_split[1]} {nome_split[2]} {nome_split[3]}")
	login = (f"{nome_split[0]}.{nome_split[3]}")
if len(nome_split) == 5:
	Nome = (f"{nome_split[0]} {nome_split[1]} {nome_split[2]} {nome_split[3]} {nome_split[4]}")
	login = (f"{nome_split[0]}.{nome_split[4]}")
if len(nome_split) == 6:
	Nome = (f"{nome_split[0]} {nome_split[1]} {nome_split[2]} {nome_split[3]} {nome_split[4]} {nome_split[5]}")
	login = (f"{nome_split[0]}.{nome_split[5]}")
if len(nome_split) == 7:
	Nome = (f"{nome_split[0]} {nome_split[1]} {nome_split[2]} {nome_split[3]} {nome_split[4]} {nome_split[5]} {nome_split[6]}")
	login = (f"{nome_split[0]}.{nome_split[6]}")

#Abrindo arquivo que vai ser editado
documento = Document("Admissao.docx")

#Trocando caracteres no texto do word
for paragrafo in documento.paragraphs:
	paragrafo.text = paragrafo.text.replace("nome_no_texto" , Nome)
	paragrafo.text = paragrafo.text.replace("usu_login" , login.lower())
	paragrafo.text = paragrafo.text.replace("usu_cpf" , cpf)
	paragrafo.text = paragrafo.text.replace("usu_departamento" , departamento)

#Salvando conteudo em um novo arquivo word	
documento.save(f"CH {NumChamado} - {Nome} - ADMISSAO.docx")
